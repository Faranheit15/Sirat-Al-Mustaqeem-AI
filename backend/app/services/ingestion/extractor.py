from __future__ import annotations

import logging
from dataclasses import dataclass
from io import BytesIO

import chardet
import pypdf

logger = logging.getLogger(__name__)

try:
    import pytesseract
    from PIL import Image

    _OCR_AVAILABLE = True
except ImportError:
    _OCR_AVAILABLE = False
    logger.debug("pytesseract/Pillow not available — OCR disabled")

_MIN_CHARS_PER_PAGE = 50


@dataclass
class ExtractedDocument:
    text: str
    language_detected: str | None
    page_count: int
    is_ocr: bool


def _detect_language(text: str) -> str | None:
    if not text.strip():
        return None
    arabic = sum(1 for c in text if "؀" <= c <= "ۿ")
    ratio = arabic / max(len(text.strip()), 1)
    if ratio > 0.2:
        return "ar"
    return "en"


def _ocr_pdf_pages(reader: pypdf.PdfReader) -> str:
    parts: list[str] = []
    for page in reader.pages:
        for img_obj in page.images:
            try:
                img = Image.open(BytesIO(img_obj.data))
                text = pytesseract.image_to_string(img, lang="ara+urd+eng")
                parts.append(text)
            except Exception as exc:
                logger.debug("ocr_page_image_failed | error=%s", exc)
    return "\n".join(parts).strip()


def extract_pdf(file_bytes: bytes) -> ExtractedDocument:
    reader = pypdf.PdfReader(BytesIO(file_bytes))
    page_count = len(reader.pages)
    text_parts = [page.extract_text() or "" for page in reader.pages]
    text = "\n".join(text_parts).strip()

    if len(text) < _MIN_CHARS_PER_PAGE * page_count and _OCR_AVAILABLE:
        logger.debug("pdf_text_sparse | pages=%d chars=%d — attempting OCR", page_count, len(text))
        ocr_text = _ocr_pdf_pages(reader)
        if ocr_text:
            return ExtractedDocument(
                text=ocr_text,
                language_detected=_detect_language(ocr_text),
                page_count=page_count,
                is_ocr=True,
            )

    return ExtractedDocument(
        text=text,
        language_detected=_detect_language(text),
        page_count=page_count,
        is_ocr=False,
    )


def extract_docx(file_bytes: bytes) -> ExtractedDocument:
    from docx import Document as DocxDocument

    doc = DocxDocument(BytesIO(file_bytes))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    text = "\n".join(paragraphs)
    page_estimate = max(len(paragraphs) // 30, 1)
    return ExtractedDocument(
        text=text,
        language_detected=_detect_language(text),
        page_count=page_estimate,
        is_ocr=False,
    )


def extract_txt(file_bytes: bytes) -> ExtractedDocument:
    detected = chardet.detect(file_bytes)
    encoding: str = detected.get("encoding") or "utf-8"
    try:
        text = file_bytes.decode(encoding)
    except (UnicodeDecodeError, LookupError):
        text = file_bytes.decode("utf-8", errors="replace")
    return ExtractedDocument(
        text=text,
        language_detected=_detect_language(text),
        page_count=1,
        is_ocr=False,
    )


def extract_document(file_bytes: bytes, file_type: str) -> ExtractedDocument:
    fmt = file_type.lower().lstrip(".")
    if fmt == "pdf":
        return extract_pdf(file_bytes)
    if fmt in ("docx", "doc"):
        return extract_docx(file_bytes)
    if fmt == "txt":
        return extract_txt(file_bytes)
    raise ValueError(f"Unsupported file type: {file_type!r}")
